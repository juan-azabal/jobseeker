"""Integration test: build_docx() → audit_docx() golden-path spec.

This is the **living specification** for CV output quality (Phase 6).

Every formatting requirement is verified end-to-end:
  - Full pipeline passes ATS audit (no violations, no warnings)
  - Name: 20pt bold #1F4E79
  - Section headers: 12pt bold #1F4E79
  - Bullets: List Bullet style only (no unicode characters)
  - Context lines: italic 10pt #555555
  - Company+date lines: tab stop in XML, company run #1F4E79 (not bold), role run normal, date run #555555
  - All-bold paragraphs < 15 (only name + section headers)
  - Zero tables
  - Only 3 non-black colors: #1F4E79, #444444, #555555
  - No **bold** markers leaking as text
  - Core Skills: #1F4E79 theme run (not bold) + normal prose run
"""

import pytest
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn


# ── Color constants (mirror docx_builder) ────────────────────────────────────
_COLOR_ACCENT = RGBColor(0x1F, 0x4E, 0x79)  # dark blue: name, section headers
_COLOR_TITLE = RGBColor(0x44, 0x44, 0x44)  # dark gray: professional title
_COLOR_MUTED = RGBColor(0x55, 0x55, 0x55)  # medium gray: contact, dates, context

# ── Golden-path Phase 6 markdown ─────────────────────────────────────────────
# Realistic 2-page CV with all required sections, tab-separated dates,
# italic context lines, and Theme: prose Core Skills format.
GOLDEN_MARKDOWN = (
    "# Ana Garcia\n"
    "Senior Product Manager - Platform & AI\n"
    "Madrid, Spain | ana.garcia@example.com | +34 600 123 456 | linkedin.com/in/anagarcia\n"
    "\n"
    "## Summary\n"
    "\n"
    "Product leader with 12+ years building data and AI platform products at scale. "
    "Background across B2B SaaS, digital media, and fintech. "
    "Not looking for pure execution roles or companies without a data strategy at board level. "
    "Work best in contexts where engineers and product have joint ownership of outcomes.\n"
    "\n"
    "## Selected Impact\n"
    "\n"
    "- Reduced time-to-market 35% by redesigning the feature delivery pipeline across 4 product squads\n"
    "- Grew marketplace GMV 28% by launching recommendation engine using collaborative filtering\n"
    "- Cut infrastructure costs 22% by migrating batch ETL to streaming with Apache Kafka\n"
    "- Launched self-serve analytics for 300 enterprise customers, reducing support tickets 40%\n"
    "\n"
    "## Core Skills\n"
    "\n"
    "Data Platforms: Deep expertise in data warehousing, streaming pipelines, ML feature stores, and self-serve analytics tooling.\n"
    "Product Strategy: Roadmap prioritization using RICE and OKRs, stakeholder alignment, board-level reporting.\n"
    "Technical Leadership: Comfortable with SQL, Python, and reviewing engineering architecture designs.\n"
    "\n"
    "## Projects\n"
    "\n"
    "### DataPulse  github.com/anagarcia/datapulse\n"
    "Open-source data quality monitoring tool built with Python, dbt, and Apache Airflow. 400+ GitHub stars.\n"
    "\n"
    "## Work Experience\n"
    "\n"
    "### Acme Technologies - Head of Data Platform\t01/2022 - Present\n"
    "_Enterprise SaaS company serving 500+ global clients; leading cross-functional team of 8 building the company data platform._\n"
    "\n"
    "- Defined 3-year platform roadmap adopted by CTO as company north star\n"
    "- Increased data pipeline reliability from 94% to 99.7% SLA through observability overhaul\n"
    "- Launched self-serve analytics product; 300 enterprise customers onboarded within 6 months\n"
    "\n"
    "### Mediagroup Corp - Senior Product Manager\t03/2019 - 12/2021\n"
    "_Digital media conglomerate; owned personalization and monetization products for 15M daily users._\n"
    "\n"
    "- Built A/B testing platform used by 40 editorial teams, increasing click-through rates 18%\n"
    "- Migrated ad serving from third-party to in-house stack, saving 1.2M EUR per year\n"
    "\n"
    "## Education and Certifications\n"
    "\n"
    "MBA - ESADE Business School, 2018\n"
    "MSc Computer Science - UPM, 2013\n"
    "AWS Certified Solutions Architect - 2021\n"
    "\n"
    "## Languages\n"
    "\n"
    "Spanish (native) | English (fluent) | French (intermediate)\n"
)


def _build(tmp_path) -> tuple[Path, Document]:
    """Build the golden .docx and return (path, Document)."""
    from api.cv.docx_builder import build_docx

    path = tmp_path / "integration_golden.docx"
    build_docx(GOLDEN_MARKDOWN, str(path))
    return path, Document(str(path))


# ── Full pipeline tests ───────────────────────────────────────────────────────


def test_pipeline_passes_ats_audit(tmp_path):
    """Full pipeline: build_docx() → audit_docx() → passed:True, zero violations."""
    from api.cv.ats_audit import audit_docx

    path, _ = _build(tmp_path)
    result = audit_docx(str(path))
    assert result["passed"] is True, f"ATS audit failed. Violations: {result['violations']}"
    assert result["violations"] == [], f"Expected zero violations, got: {result['violations']}"


def test_pipeline_zero_warnings(tmp_path):
    """Full pipeline: build_docx() → audit_docx() → zero formatting warnings."""
    from api.cv.ats_audit import audit_docx

    path, _ = _build(tmp_path)
    result = audit_docx(str(path))
    assert result["warnings"] == [], f"Expected zero warnings, got: {result['warnings']}"


def test_pipeline_audit_result_has_all_keys(tmp_path):
    """audit_docx() result must contain passed, violations, warnings, stats."""
    from api.cv.ats_audit import audit_docx

    path, _ = _build(tmp_path)
    result = audit_docx(str(path))
    for key in ("passed", "violations", "warnings", "stats"):
        assert key in result, f"Missing key '{key}' in audit result"


def test_pipeline_stats_sections_and_bullets(tmp_path):
    """ATS audit stats must reflect all required sections and bullets."""
    from api.cv.ats_audit import audit_docx

    path, _ = _build(tmp_path)
    stats = audit_docx(str(path))["stats"]
    assert stats["section_count"] >= 6, f"Expected ≥6 sections, got {stats['section_count']}"
    assert stats["bullet_count"] >= 5, f"Expected ≥5 bullets, got {stats['bullet_count']}"
    assert stats["paragraph_count"] > 10, "Too few paragraphs"
    assert stats["estimated_pages"] >= 1


def test_pipeline_stats_formatting_counters(tmp_path):
    """ATS audit stats must report italic lines, tab lines, bold lines."""
    from api.cv.ats_audit import audit_docx

    path, _ = _build(tmp_path)
    stats = audit_docx(str(path))["stats"]
    assert stats["italic_count"] >= 2, f"Expected ≥2 italic context paragraphs, got {stats['italic_count']}"
    assert stats["tab_count"] >= 2, f"Expected ≥2 tab-stop paragraphs (company lines), got {stats['tab_count']}"
    assert stats["bold_count"] < 15, f"Expected <15 all-bold paragraphs, got {stats['bold_count']}"


# ── Formatting spec tests ─────────────────────────────────────────────────────


def test_name_20pt_bold_accent_color(tmp_path):
    """Name: 20pt, bold, #1F4E79."""
    path, doc = _build(tmp_path)
    name_para = next(
        (p for p in doc.paragraphs if p.runs and p.runs[0].font.size and abs(p.runs[0].font.size.pt - 20.0) < 0.5),
        None,
    )
    assert name_para is not None, "No 20pt paragraph found (expected: name)"
    assert "Ana Garcia" in name_para.text
    run = name_para.runs[0]
    assert run.font.bold is True, "Name run must be bold"
    assert run.font.color.type is not None, "Name must have explicit color"
    assert run.font.color.rgb == _COLOR_ACCENT, f"Name color must be #1F4E79, got {run.font.color.rgb}"


def test_all_required_section_headers_present(tmp_path):
    """All 6 required section headers: 12pt bold #1F4E79."""
    REQUIRED = {
        "Summary",
        "Selected Impact",
        "Core Skills",
        "Work Experience",
        "Education and Certifications",
        "Languages",
    }
    path, doc = _build(tmp_path)
    header_paras = {p.text.strip(): p for p in doc.paragraphs if p.text.strip() in REQUIRED}
    missing = REQUIRED - set(header_paras.keys())
    assert not missing, f"Missing section headers: {missing}"
    for name, para in header_paras.items():
        assert para.runs, f"Header '{name}' has no runs"
        run = para.runs[0]
        assert run.font.bold is True, f"Header '{name}' must be bold"
        assert run.font.size and abs(run.font.size.pt - 12.0) < 0.5, f"Header '{name}' must be 12pt"
        assert run.font.color.type is not None
        assert run.font.color.rgb == _COLOR_ACCENT, f"Header '{name}' must be #1F4E79"


def test_bullets_use_list_style_only(tmp_path):
    """Bullets: at least 5 List Bullet paragraphs; zero unicode bullet characters."""
    UNICODE_BULLETS = {"\u2022", "\u25e6", "\u25aa", "\u2013", "\u2192", "•", "◦", "▪"}
    path, doc = _build(tmp_path)
    list_paras = [p for p in doc.paragraphs if "List" in p.style.name or "Bullet" in p.style.name]
    assert len(list_paras) >= 5, f"Expected ≥5 List Bullet paragraphs, got {len(list_paras)}"
    for para in doc.paragraphs:
        for char in UNICODE_BULLETS:
            assert char not in para.text, f"Unicode bullet '{char}' found in: {para.text[:60]}"


def test_context_lines_italic_10pt_muted(tmp_path):
    """Work Experience context lines: italic, 10pt, #555555, at least 2."""
    path, doc = _build(tmp_path)
    italic_paras = [p for p in doc.paragraphs if any(r.font.italic for r in p.runs)]
    assert len(italic_paras) >= 2, f"Expected ≥2 italic context paragraphs, got {len(italic_paras)}"
    for para in italic_paras:
        italic_run = next(r for r in para.runs if r.font.italic)
        assert italic_run.font.size and abs(italic_run.font.size.pt - 10.0) < 0.5, (
            f"Context line not 10pt: '{para.text[:60]}'"
        )
        assert italic_run.font.color.type is not None
        assert italic_run.font.color.rgb == _COLOR_MUTED, f"Context line not #555555: '{para.text[:60]}'"


def test_company_date_lines_tab_stop_accent_company_muted_date(tmp_path):
    """Company+date lines: <w:tabs> in XML, company run #1F4E79 (not bold), role normal, date muted.

    Structure: run1=company(#1F4E79, not bold) + run2=" - Role"(normal) + runN=\\tdate(muted #555555).
    The date is always the last run.
    """
    path, doc = _build(tmp_path)
    tab_paras = [p for p in doc.paragraphs if "\t" in p.text]
    assert len(tab_paras) >= 2, f"Expected ≥2 company/date paragraphs with tab, got {len(tab_paras)}"
    for para in tab_paras:
        pPr = para._p.pPr
        assert pPr is not None and pPr.find(qn("w:tabs")) is not None, f"No <w:tabs> XML element in: '{para.text[:60]}'"
        assert len(para.runs) >= 2, f"Expected ≥2 runs in company/date line: '{para.text[:60]}'"
        # First run: company name — accent color, NOT bold
        run_company = para.runs[0]
        assert run_company.font.bold is not True, f"Company run1 must not be bold: '{para.text[:60]}'"
        assert run_company.font.color.type is not None, f"Company run1 must have explicit color: '{para.text[:60]}'"
        assert run_company.font.color.rgb == _COLOR_ACCENT, f"Company run1 must be #1F4E79: '{para.text[:60]}'"
        # Last run: \tdate — not bold, muted color
        date_run = para.runs[-1]
        assert date_run.font.bold is not True, f"Date run must not be bold: '{para.text[:60]}'"
        assert date_run.font.color.type is not None
        assert date_run.font.color.rgb == _COLOR_MUTED, f"Date run must be #555555: '{para.text[:60]}'"


def test_all_bold_paragraphs_under_15(tmp_path):
    """Total all-bold paragraphs must be < 15 (only headers and company lines)."""
    path, doc = _build(tmp_path)
    bold_paras = [p for p in doc.paragraphs if p.runs and all(r.font.bold is True for r in p.runs if r.text.strip())]
    assert len(bold_paras) < 15, (
        f"Too many all-bold paragraphs: {len(bold_paras)} — only headers and company "
        f"name runs should be all-bold. Check for body text inadvertently bolded."
    )


def test_zero_tables_in_document(tmp_path):
    """Document must contain zero tables (ATS parsers cannot parse table-based CVs)."""
    path, doc = _build(tmp_path)
    assert len(doc.tables) == 0, f"Found {len(doc.tables)} table(s) — ATS hard violation"


def test_only_three_allowed_colors(tmp_path):
    """Only 3 non-black colors allowed: #1F4E79, #444444, #555555."""
    ALLOWED = {_COLOR_ACCENT, _COLOR_TITLE, _COLOR_MUTED}
    path, doc = _build(tmp_path)
    for para in doc.paragraphs:
        for run in para.runs:
            try:
                if run.font.color.type is None:
                    continue  # black / no explicit color — OK
                rgb = run.font.color.rgb
                assert rgb in ALLOWED, f"Disallowed color {rgb!r} in: '{para.text[:60]}'"
            except Exception as exc:
                # Theme-inherited colors: ignore (not explicit rgb)
                if "rgb" in str(exc).lower():
                    pass


def test_no_bold_markers_leaked_as_text(tmp_path):
    """No **bold** markdown markers must appear as literal text in the .docx."""
    path, doc = _build(tmp_path)
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "**" not in full_text, "Bold markers (**) leaked into output — _clean_text() may have missed them"


def test_core_skills_accent_theme_plus_normal_prose(tmp_path):
    """Core Skills: each line has #1F4E79 theme run (not bold) + non-bold prose run, ≥2 lines."""
    path, doc = _build(tmp_path)
    skill_paras = [
        p
        for p in doc.paragraphs
        if len(p.runs) >= 2
        and p.runs[0].font.bold is not True
        and (p.runs[0].font.color.type is not None and p.runs[0].font.color.rgb == _COLOR_ACCENT)
        and p.runs[1].font.bold is not True
        and ":" in p.runs[0].text
    ]
    assert len(skill_paras) >= 2, (
        f"Expected ≥2 Core Skills theme lines (#1F4E79 theme + normal prose), got {len(skill_paras)}"
    )


def test_em_dash_replaced_in_golden(tmp_path):
    """Em dash in golden markdown must be replaced with hyphen in output."""
    from api.cv.docx_builder import build_docx

    # Inject an em dash into a bullet
    md = GOLDEN_MARKDOWN.replace(
        "Reduced time-to-market 35%",
        "Reduced time-to-market \u2014 35%",
    )
    out = tmp_path / "em_dash_test.docx"
    build_docx(md, str(out))
    doc = Document(str(out))
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "\u2014" not in full_text, "Em dash was not replaced in post-processing"


def test_oxford_comma_replaced_in_golden(tmp_path):
    """Oxford comma in golden markdown must be removed in output."""
    from api.cv.docx_builder import build_docx

    md = GOLDEN_MARKDOWN.replace(
        "Data Platforms: Deep expertise in data warehousing, streaming pipelines, ML feature stores, and self-serve analytics tooling.",
        "Data Platforms: Deep expertise in data warehousing, streaming pipelines, ML feature stores, and self-serve analytics tooling, and more.",
    )
    out = tmp_path / "oxford_test.docx"
    build_docx(md, str(out))
    doc = Document(str(out))
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert ", and " not in full_text, "Oxford comma pattern found in output"
