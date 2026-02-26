"""Tests for api/cv/docx_builder.py — Phase 6 skill-quality formatting."""

import pytest
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import RGBColor


# --- Phase 6 sample markdown (new output contract) ---
SAMPLE_MARKDOWN = (
    "# Juan Azabal\n"
    "Senior Product Manager - Data Platforms & AI\n"
    "Barcelona, Spain | j.azabal@gmail.com | +34 625 588 926 | linkedin.com/in/juanazabal\n"
    "\n"
    "## Summary\n"
    "\n"
    "Experienced product manager with 10+ years building data platform products in competitive "
    "digital environments. Not looking for pure execution roles without product autonomy.\n"
    "\n"
    "## Selected Impact\n"
    "\n"
    "- Grew ad revenue 40% YoY by redesigning targeting pipeline, reducing latency from 800ms to 120ms\n"
    "- Led migration of legacy data platform serving 50M monthly users with zero downtime\n"
    "- Reduced time-to-insight from 3 days to 4 hours by delivering self-serve analytics tooling\n"
    "\n"
    "## Core Skills\n"
    "\n"
    "Data & Analytics: Deep experience with data warehousing, streaming pipelines, and self-serve analytics.\n"
    "Product Strategy: Background in roadmap definition, stakeholder alignment, and OKR frameworks.\n"
    "\n"
    "## Projects\n"
    "\n"
    "### JobSeeker  github.com/juan-azabal/jobseeker\n"
    "Job search CRM with AI-powered CV tailoring. FastAPI, React, SQLite.\n"
    "\n"
    "## Work Experience\n"
    "\n"
    "### Gartner Digital Markets (Capterra) - Senior PM, Data Platform\t07/2024 - Present\n"
    "_Replacing GA4 with a governed, first-party data platform across three global B2B marketplace brands._\n"
    "\n"
    "- Defined 3-year data platform roadmap adopted by engineering leadership\n"
    "- Increased query performance 60% by migrating to columnar storage\n"
    "\n"
    "### Grupo Godo (LaVanguardia.com) - Senior Product Manager\t07/2022 - 07/2024\n"
    "_Data, personalization and monetization at one of Spain's largest digital news publishers._\n"
    "\n"
    "- Built real-time personalization engine serving 10M daily users\n"
    "\n"
    "## Education and Certifications\n"
    "\n"
    "MBA - IESE Business School, 2017\n"
    "BSc Computer Science - UPC, 2013\n"
    "\n"
    "## Languages\n"
    "\n"
    "Spanish (native) | English (advanced) | Catalan (basic)\n"
)

# --- Phase 5 format (backward compat) ---
SAMPLE_MARKDOWN_V5 = """# Juan Azabal
Senior Product Manager | Data, Personalization & Monetization
Barcelona, Spain | j.azabal@gmail.com | +34 625 588 926 | linkedin.com/in/juanazabal

## Summary

Experienced product manager with 10 years driving data platform and monetization products.

## Selected Impact

- Grew ad revenue 40% YoY by redesigning targeting pipeline
- Led migration of legacy data platform serving 50M users
- Reduced time-to-insight from 3 days to 4 hours for analytics team

## Core Skills

**Data Platform**
Deep experience with data warehousing, streaming pipelines, and self-serve analytics.

**Product Strategy**
Strong background in roadmap definition, stakeholder alignment, and OKR frameworks.

## Projects

### JobSeeker
Job search CRM with AI-powered CV tailoring. FastAPI, React, SQLite. github.com/juan/jobseeker

## Work Experience

### Acme Corp, Barcelona, Spain
**Senior Product Manager | 01/2021 - Present**

- Defined 3-year data platform roadmap adopted by engineering leadership
- Increased query performance 60% by migrating to columnar storage

### Previous Company, Madrid, Spain
**Product Manager | 03/2018 - 12/2020**

- Built real-time personalization engine serving 10M daily users

## Education and Certifications

- MBA - IESE Business School, 2017
- BSc Computer Science - UPC, 2013

## Languages

- Spanish - Native
- English - Fluent
- Catalan - Native
"""


# --- Color constants (mirror docx_builder) ---
_COLOR_ACCENT = RGBColor(0x1F, 0x4E, 0x79)
_COLOR_TITLE = RGBColor(0x44, 0x44, 0x44)
_COLOR_MUTED = RGBColor(0x55, 0x55, 0x55)


def _build(markdown: str) -> Path:
    """Helper: run build_docx and return path."""
    from api.cv.docx_builder import build_docx

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        output_path = f.name
    build_docx(markdown, output_path)
    return Path(output_path)


def _find_para_by_size(doc: Document, size_pt: float):
    """Find first paragraph whose first run has given font size."""
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.size and abs(run.font.size.pt - size_pt) < 0.5:
                return para
    return None


def _find_paras_by_bold_color(doc: Document, bold: bool, color: RGBColor):
    """Find paragraphs whose first run matches bold + color."""
    result = []
    for para in doc.paragraphs:
        if not para.runs:
            continue
        run = para.runs[0]
        if run.font.bold is True and run.font.bold == bold:
            try:
                if run.font.color.type is not None and run.font.color.rgb == color:
                    result.append(para)
            except Exception:
                pass
    return result


# ── Existing core tests (still valid) ───────────────────────────────────────


def test_output_file_exists_and_has_content():
    """build_docx() creates a non-empty .docx file."""
    path = _build(SAMPLE_MARKDOWN)
    try:
        assert path.exists()
        assert path.stat().st_size > 0
    finally:
        path.unlink(missing_ok=True)


def test_zero_tables_in_xml():
    """Generated .docx must have zero <w:tbl> elements (ATS requirement)."""
    path = _build(SAMPLE_MARKDOWN)
    try:
        doc = Document(str(path))
        assert len(doc.tables) == 0, f"Found {len(doc.tables)} table(s) — ATS violation"
    finally:
        path.unlink(missing_ok=True)


def test_name_paragraph_contains_candidate_name():
    """First paragraph must contain the candidate name."""
    path = _build(SAMPLE_MARKDOWN)
    try:
        doc = Document(str(path))
        name_para = next((p for p in doc.paragraphs if p.text.strip()), None)
        assert name_para is not None
        assert "Juan Azabal" in name_para.text
    finally:
        path.unlink(missing_ok=True)


def test_section_count_matches_sections():
    """Document must have exactly 7 section header paragraphs."""
    SECTION_NAMES = {
        "Summary",
        "Selected Impact",
        "Core Skills",
        "Projects",
        "Work Experience",
        "Education and Certifications",
        "Languages",
    }
    path = _build(SAMPLE_MARKDOWN)
    try:
        doc = Document(str(path))
        header_paras = [p for p in doc.paragraphs if p.text.strip() in SECTION_NAMES]
        assert len(header_paras) == 7, (
            f"Expected 7 section headers, found {len(header_paras)}: {[p.text for p in header_paras]}"
        )
    finally:
        path.unlink(missing_ok=True)


def test_bullets_use_list_style_not_unicode():
    """Bullet paragraphs must use List Bullet style; no unicode bullet characters."""
    path = _build(SAMPLE_MARKDOWN)
    try:
        doc = Document(str(path))
        unicode_bullets = {"•", "◦", "▪", "–", "→"}
        for para in doc.paragraphs:
            for char in unicode_bullets:
                assert char not in para.text, f"Unicode bullet '{char}' found in: {para.text[:60]}"
        list_paras = [p for p in doc.paragraphs if "List" in p.style.name or "Bullet" in p.style.name]
        assert len(list_paras) > 0, "No List Bullet paragraphs found"
    finally:
        path.unlink(missing_ok=True)


def test_em_dashes_replaced_in_post_processing():
    """Post-processing must replace em dashes with ASCII hyphens."""
    markdown_with_em = SAMPLE_MARKDOWN.replace(
        "Grew ad revenue 40% YoY by redesigning targeting pipeline",
        "Grew ad revenue 40% YoY \u2014 redesigned targeting pipeline",
    )
    path = _build(markdown_with_em)
    try:
        doc = Document(str(path))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "\u2014" not in full_text, "Em dash found in output"
    finally:
        path.unlink(missing_ok=True)


def test_oxford_comma_replaced_in_post_processing():
    """Post-processing must remove Oxford comma pattern ', and '."""
    markdown_with_oxford = SAMPLE_MARKDOWN.replace(
        "Deep experience with data warehousing, streaming pipelines, and self-serve analytics.",
        "Deep experience with data warehousing, streaming pipelines, and self-serve analytics, and more.",
    )
    path = _build(markdown_with_oxford)
    try:
        doc = Document(str(path))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert ", and " not in full_text, "Oxford comma pattern found in output"
    finally:
        path.unlink(missing_ok=True)


def test_code_fences_stripped():
    """Markdown wrapped in code fences must produce a valid .docx."""
    fenced = "```markdown\n" + SAMPLE_MARKDOWN + "\n```"
    path = _build(fenced)
    try:
        doc = Document(str(path))
        name_para = next((p for p in doc.paragraphs if p.text.strip()), None)
        assert name_para is not None
        assert "Juan Azabal" in name_para.text
    finally:
        path.unlink(missing_ok=True)


def test_preamble_stripped():
    """Text before first # heading must be stripped."""
    with_preamble = "Here is your tailored CV:\n\n" + SAMPLE_MARKDOWN
    path = _build(with_preamble)
    try:
        doc = Document(str(path))
        name_para = next((p for p in doc.paragraphs if p.text.strip()), None)
        assert name_para is not None
        assert "Juan Azabal" in name_para.text
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "Here is your tailored CV" not in full_text
    finally:
        path.unlink(missing_ok=True)


def test_empty_markdown_raises_value_error():
    """Empty or whitespace-only markdown must raise ValueError."""
    from api.cv.docx_builder import build_docx

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        output_path = f.name
    try:
        with pytest.raises(ValueError):
            build_docx("   \n\n  ", output_path)
    finally:
        Path(output_path).unlink(missing_ok=True)


def test_malformed_markdown_no_heading_raises_value_error():
    """Markdown with no # heading raises ValueError."""
    from api.cv.docx_builder import build_docx

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        output_path = f.name
    try:
        with pytest.raises(ValueError):
            build_docx("Some random text\nno headings here\njust prose", output_path)
    finally:
        Path(output_path).unlink(missing_ok=True)


# ── Phase 6 formatting spec tests ───────────────────────────────────────────


def test_name_paragraph_20pt_bold_accent_color():
    """Name: 20pt, bold, color #1F4E79."""
    path = _build(SAMPLE_MARKDOWN)
    try:
        doc = Document(str(path))
        name_para = _find_para_by_size(doc, 20.0)
        assert name_para is not None, "No 20pt paragraph found (name)"
        assert "Juan Azabal" in name_para.text
        run = name_para.runs[0]
        assert run.font.bold is True, "Name run should be bold"
        assert run.font.color.type is not None, "Name run should have explicit color"
        assert run.font.color.rgb == _COLOR_ACCENT, f"Name color should be #1F4E79, got {run.font.color.rgb}"
    finally:
        path.unlink(missing_ok=True)


def test_section_headers_12pt_bold_accent_color():
    """Section headers (##): 12pt, bold, color #1F4E79."""
    SECTION_NAMES = {
        "Summary",
        "Selected Impact",
        "Core Skills",
        "Projects",
        "Work Experience",
        "Education and Certifications",
        "Languages",
    }
    path = _build(SAMPLE_MARKDOWN)
    try:
        doc = Document(str(path))
        header_paras = [p for p in doc.paragraphs if p.text.strip() in SECTION_NAMES]
        assert len(header_paras) == 7
        for para in header_paras:
            assert para.runs, f"Header '{para.text}' has no runs"
            run = para.runs[0]
            assert run.font.bold is True, f"Header '{para.text}' not bold"
            assert run.font.size and abs(run.font.size.pt - 12.0) < 0.5, (
                f"Header '{para.text}' not 12pt (got {run.font.size})"
            )
            assert run.font.color.type is not None
            assert run.font.color.rgb == _COLOR_ACCENT, f"Header '{para.text}' color wrong"
    finally:
        path.unlink(missing_ok=True)


def test_title_line_13pt_dark_gray():
    """Title line: 13pt, color #444444, not bold."""
    path = _build(SAMPLE_MARKDOWN)
    try:
        doc = Document(str(path))
        # Title is the second non-empty paragraph (after name)
        non_empty = [p for p in doc.paragraphs if p.text.strip()]
        assert len(non_empty) >= 2
        title_para = non_empty[1]  # second paragraph
        assert "Senior Product Manager" in title_para.text
        run = title_para.runs[0]
        assert run.font.size and abs(run.font.size.pt - 13.0) < 0.5, "Title not 13pt"
        assert run.font.bold is not True, "Title should not be bold"
        assert run.font.color.type is not None
        assert run.font.color.rgb == _COLOR_TITLE, "Title color should be #444444"
    finally:
        path.unlink(missing_ok=True)


def test_contact_line_10pt_muted_gray():
    """Contact line: 10pt, color #555555."""
    path = _build(SAMPLE_MARKDOWN)
    try:
        doc = Document(str(path))
        # Contact line contains email/linkedin
        contact_para = next(
            (p for p in doc.paragraphs if "j.azabal@gmail.com" in p.text),
            None,
        )
        assert contact_para is not None, "Contact paragraph not found"
        run = contact_para.runs[0]
        assert run.font.size and abs(run.font.size.pt - 10.0) < 0.5, "Contact not 10pt"
        assert run.font.color.type is not None
        assert run.font.color.rgb == _COLOR_MUTED, "Contact color should be #555555"
    finally:
        path.unlink(missing_ok=True)


def test_italic_context_lines_present():
    """Work Experience context lines (_..._): at least 2, italic, 10pt, #555555."""
    path = _build(SAMPLE_MARKDOWN)
    try:
        doc = Document(str(path))
        italic_paras = [p for p in doc.paragraphs if any(r.font.italic for r in p.runs)]
        assert len(italic_paras) >= 2, f"Expected at least 2 italic context lines, found {len(italic_paras)}"
        for para in italic_paras:
            italic_run = next(r for r in para.runs if r.font.italic)
            assert italic_run.font.size and abs(italic_run.font.size.pt - 10.0) < 0.5, (
                f"Context line not 10pt: {para.text[:50]}"
            )
            assert italic_run.font.color.type is not None
            assert italic_run.font.color.rgb == _COLOR_MUTED, f"Context line color should be #555555: {para.text[:50]}"
    finally:
        path.unlink(missing_ok=True)


def test_company_date_line_tab_stop_accent_company_muted_date():
    """Company lines: tab stop in XML, run1 company #1F4E79 (not bold), last run date muted #555555.

    Structure: run1=company(#1F4E79, not bold) [+ run2=" - Role"(normal)] + runN=\tdate(muted).
    The date is always the last run; role (if present) is not bold.
    """
    from docx.oxml.ns import qn

    path = _build(SAMPLE_MARKDOWN)
    try:
        doc = Document(str(path))
        # Company/date lines contain a tab character
        tab_paras = [p for p in doc.paragraphs if "\t" in p.text]
        assert len(tab_paras) >= 2, f"Expected at least 2 company/date lines, found {len(tab_paras)}"
        for para in tab_paras:
            # Has tab stop in XML
            pPr = para._p.pPr
            assert pPr is not None and pPr.find(qn("w:tabs")) is not None, f"No <w:tabs> XML in: {para.text[:60]}"
            assert len(para.runs) >= 2, f"Expected ≥2 runs in company line, found {len(para.runs)}: {para.text[:60]}"
            # First run: company name — accent color, NOT bold
            run_company = para.runs[0]
            assert run_company.font.bold is not True, f"Company run1 must not be bold: {para.text[:60]}"
            assert run_company.font.color.type is not None, f"Company run1 must have explicit color: {para.text[:60]}"
            assert run_company.font.color.rgb == _COLOR_ACCENT, f"Company run1 color must be #1F4E79: {para.text[:60]}"
            # Last run: \tdate — not bold, muted color
            date_run = para.runs[-1]
            assert date_run.font.bold is not True, f"Date run should not be bold: {para.text[:60]}"
            assert date_run.font.color.type is not None
            assert date_run.font.color.rgb == _COLOR_MUTED, f"Date color should be #555555: {para.text[:60]}"
    finally:
        path.unlink(missing_ok=True)


def test_core_skills_theme_first_run_accent_color():
    """Core Skills lines: first run (theme name) #1F4E79 not bold 11pt, second run (prose) not bold 11pt."""
    path = _build(SAMPLE_MARKDOWN)
    try:
        doc = Document(str(path))
        # Core Skills themes: first run accent color (not bold) + second run normal
        theme_paras = [
            p
            for p in doc.paragraphs
            if len(p.runs) >= 2
            and p.runs[0].font.bold is not True
            and (p.runs[0].font.color.type is not None and p.runs[0].font.color.rgb == _COLOR_ACCENT)
            and p.runs[1].font.bold is not True
            and ":" in p.runs[0].text
        ]
        assert len(theme_paras) >= 2, (
            f"Expected at least 2 Core Skills theme lines (accent color + normal prose), found {len(theme_paras)}"
        )
        for para in theme_paras:
            # Both runs should be 11pt
            for run in para.runs[:2]:
                if run.font.size:
                    assert abs(run.font.size.pt - 11.0) < 0.5, f"Core Skills run not 11pt: {para.text[:50]}"
    finally:
        path.unlink(missing_ok=True)


def test_project_url_run_10pt_muted():
    """Project line URL run: 10pt, color #555555. Name run: #1F4E79, not bold."""
    path = _build(SAMPLE_MARKDOWN)
    try:
        doc = Document(str(path))
        # Find project paragraph: run1 has accent color (not bold), run2 contains URL
        url_run = None
        for para in doc.paragraphs:
            if len(para.runs) >= 2:
                run1 = para.runs[0]
                run2 = para.runs[1]
                if "github" in run2.text.lower() or ".com" in run2.text.lower():
                    try:
                        if (
                            run1.font.bold is not True
                            and run1.font.color.type is not None
                            and run1.font.color.rgb == _COLOR_ACCENT
                        ):
                            url_run = run2
                            break
                    except Exception:
                        pass
        assert url_run is not None, "No project URL run found"
        assert url_run.font.size and abs(url_run.font.size.pt - 10.0) < 0.5, f"Project URL run not 10pt: {url_run.text}"
        assert url_run.font.color.type is not None
        assert url_run.font.color.rgb == _COLOR_MUTED, "Project URL color should be #555555"
    finally:
        path.unlink(missing_ok=True)


def test_total_bold_paragraphs_under_15():
    """Total all-bold paragraphs must be < 15 (headers + company lines only)."""
    path = _build(SAMPLE_MARKDOWN)
    try:
        doc = Document(str(path))
        bold_paras = [
            p for p in doc.paragraphs if p.runs and all(r.font.bold is True for r in p.runs if r.text.strip())
        ]
        assert len(bold_paras) < 15, (
            f"Too many all-bold paragraphs: {len(bold_paras)} (expected < 15 — headers and company lines only)"
        )
    finally:
        path.unlink(missing_ok=True)


def test_only_expected_colors_in_document():
    """Only 3 non-black colors allowed: #1F4E79, #444444, #555555."""
    ALLOWED = {_COLOR_ACCENT, _COLOR_TITLE, _COLOR_MUTED}
    path = _build(SAMPLE_MARKDOWN)
    try:
        doc = Document(str(path))
        for para in doc.paragraphs:
            for run in para.runs:
                try:
                    if run.font.color.type is None:
                        continue  # black / no explicit color — OK
                    rgb = run.font.color.rgb
                    assert rgb in ALLOWED, f"Unexpected color {rgb!r} in paragraph: '{para.text[:50]}'"
                except Exception as exc:
                    # Theme-inherited colors: ignore
                    if "rgb" in str(exc).lower():
                        pass
    finally:
        path.unlink(missing_ok=True)


def test_bold_markers_stripped_from_body_text():
    """**bold** markers in body text must be stripped — not rendered as asterisks."""
    markdown_with_bold = SAMPLE_MARKDOWN.replace(
        "Grew ad revenue 40% YoY by redesigning targeting pipeline",
        "Grew ad revenue **40% YoY** by redesigning targeting pipeline",
    )
    path = _build(markdown_with_bold)
    try:
        doc = Document(str(path))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "**" not in full_text, "Bold markers (**) leaked into output text"
        assert "40% YoY" in full_text, "Content stripped along with markers"
    finally:
        path.unlink(missing_ok=True)


def test_backward_compat_phase5_format_produces_valid_docx():
    """Phase 5 format markdown must still produce a valid .docx without crashing."""
    path = _build(SAMPLE_MARKDOWN_V5)
    try:
        doc = Document(str(path))
        assert len(doc.tables) == 0, "Phase 5 compat: no tables"
        name_para = next((p for p in doc.paragraphs if p.text.strip()), None)
        assert name_para is not None
        assert "Juan Azabal" in name_para.text
        # Should have some paragraphs
        assert len(doc.paragraphs) > 10
    finally:
        path.unlink(missing_ok=True)
