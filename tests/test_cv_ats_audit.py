"""Tests for api/cv/ats_audit.py — ATS audit utility."""

import tempfile
import pytest
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree


def _make_clean_docx(tmp_path) -> str:
    """Create a minimal ATS-compliant .docx for testing (via the real builder)."""
    from tests.test_cv_docx_builder import SAMPLE_MARKDOWN
    from api.cv.docx_builder import build_docx

    path = tmp_path / "clean.docx"
    build_docx(SAMPLE_MARKDOWN, str(path))
    return str(path)


def _make_all_headers_docx(tmp_path) -> str:
    """Create a docx with all required headers and bullets (no em dash, no table)."""
    from api.cv.docx_builder import build_docx
    from tests.test_cv_docx_builder import SAMPLE_MARKDOWN

    path = tmp_path / "all_headers.docx"
    build_docx(SAMPLE_MARKDOWN, str(path))
    return str(path)


def _make_docx_with_table(tmp_path) -> str:
    """Create a .docx that contains a table element."""
    doc = Document()
    doc.add_heading("Test Name", level=1)
    doc.add_paragraph("## Summary")
    doc.add_paragraph("A summary paragraph here.")
    doc.add_paragraph("## Work Experience")
    doc.add_paragraph("## Education and Certifications")
    doc.add_paragraph("## Languages")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Company"
    table.cell(0, 1).text = "Role"
    path = tmp_path / "with_table.docx"
    doc.save(str(path))
    return str(path)


def _make_docx_with_em_dash(tmp_path) -> str:
    """Create a .docx with an em dash in body text."""
    doc = Document()
    doc.add_heading("Test Name", level=1)
    doc.add_paragraph("Contact info here")
    h = doc.add_heading(level=2)
    h.clear()
    h.add_run("Summary")
    doc.add_paragraph("Experience \u2014 with em dash.")
    for header in ["Work Experience", "Education and Certifications", "Languages", "Selected Impact", "Core Skills"]:
        hpara = doc.add_heading(level=2)
        hpara.clear()
        hpara.add_run(header)
        doc.add_paragraph(f"Content for {header}")
    path = tmp_path / "with_em_dash.docx"
    doc.save(str(path))
    return str(path)


def _make_docx_missing_header(tmp_path) -> str:
    """Create a .docx missing the Summary standard header."""
    doc = Document()
    doc.add_heading("Test Name", level=1)
    doc.add_paragraph("Contact line")
    for h in ["Selected Impact", "Core Skills", "Work Experience", "Education and Certifications", "Languages"]:
        hpara = doc.add_heading(level=2)
        hpara.clear()
        hpara.add_run(h)
        doc.add_paragraph(f"Content for {h}")
    path = tmp_path / "missing_header.docx"
    doc.save(str(path))
    return str(path)


def _make_docx_all_bold(tmp_path) -> str:
    """Create a .docx where > 15 paragraphs are all-bold (for excessive_bold warning test).

    Uses the Phase 6 builder with a markdown that has many bold paragraphs by using
    a markdown with many section-like lines. Instead of standard sections, we create
    a raw Document with explicit bold XML to simulate the old Phase 5 all-bold issue.
    """
    doc = Document()
    # All required headers
    for header in [
        "Summary",
        "Selected Impact",
        "Core Skills",
        "Work Experience",
        "Education and Certifications",
        "Languages",
    ]:
        p = doc.add_paragraph(header)
        for run in p.runs:
            run.font.bold = True

    # Add 10 more all-bold body paragraphs (total > 15)
    for i in range(10):
        p = doc.add_paragraph(f"All bold body text paragraph {i + 1}.")
        for run in p.runs:
            run.font.bold = True

    # Add some list bullets (so no_real_bullets doesn't fire)
    doc.add_paragraph("bullet text", style="List Bullet")
    path = tmp_path / "all_bold.docx"
    doc.save(str(path))
    return str(path)


def _make_docx_no_bullets(tmp_path) -> str:
    """Create a .docx without any List Bullet paragraphs (no_real_bullets violation)."""
    doc = Document()
    # All required headers via exact text match
    for header in [
        "Summary",
        "Selected Impact",
        "Core Skills",
        "Work Experience",
        "Education and Certifications",
        "Languages",
    ]:
        p = doc.add_paragraph(header)
        run = p.runs[0]
        run.font.bold = True

    doc.add_paragraph("Prose content with no list bullets.")
    path = tmp_path / "no_bullets.docx"
    doc.save(str(path))
    return str(path)


# --- Tests ---


def test_clean_docx_passes(tmp_path):
    """A properly built ATS-compliant .docx must pass the audit."""
    from api.cv.ats_audit import audit_docx

    path = _make_clean_docx(tmp_path)
    result = audit_docx(path)
    assert result["passed"] is True
    assert result["violations"] == []


def test_audit_returns_required_keys(tmp_path):
    """audit_docx() must return passed, violations, warnings, and stats."""
    from api.cv.ats_audit import audit_docx

    path = _make_clean_docx(tmp_path)
    result = audit_docx(path)
    assert "passed" in result
    assert "violations" in result
    assert "warnings" in result
    assert "stats" in result
    assert isinstance(result["violations"], list)
    assert isinstance(result["warnings"], list)
    assert isinstance(result["stats"], dict)


def test_table_detected_as_violation(tmp_path):
    """A .docx with a table must report table_found violation."""
    from api.cv.ats_audit import audit_docx

    path = _make_docx_with_table(tmp_path)
    result = audit_docx(path)
    assert result["passed"] is False
    assert any("table" in v for v in result["violations"])


def test_em_dash_detected_as_violation(tmp_path):
    """A .docx with an em dash must report prohibited_char violation."""
    from api.cv.ats_audit import audit_docx

    path = _make_docx_with_em_dash(tmp_path)
    result = audit_docx(path)
    assert result["passed"] is False
    assert any("em_dash" in v or "prohibited_char" in v for v in result["violations"])


def test_missing_header_detected_as_violation(tmp_path):
    """A .docx missing 'Summary' header must report missing_header violation."""
    from api.cv.ats_audit import audit_docx

    path = _make_docx_missing_header(tmp_path)
    result = audit_docx(path)
    assert result["passed"] is False
    assert any("Summary" in v for v in result["violations"])


def test_stats_contain_expected_keys(tmp_path):
    """Stats dict must include section_count, bullet_count, paragraph_count, estimated_pages,
    bold_count, italic_count, tab_count."""
    from api.cv.ats_audit import audit_docx

    path = _make_clean_docx(tmp_path)
    result = audit_docx(path)
    stats = result["stats"]
    assert "section_count" in stats
    assert "bullet_count" in stats
    assert "paragraph_count" in stats
    assert "estimated_pages" in stats
    assert "bold_count" in stats
    assert "italic_count" in stats
    assert "tab_count" in stats
    assert stats["section_count"] >= 6
    assert stats["bullet_count"] > 0
    assert stats["paragraph_count"] > 0
    assert stats["estimated_pages"] >= 1


# ── Phase 6 formatting quality tests ────────────────────────────────────────


def test_no_real_bullets_is_violation(tmp_path):
    """A .docx without List Bullet paragraphs must report no_real_bullets violation."""
    from api.cv.ats_audit import audit_docx

    path = _make_docx_no_bullets(tmp_path)
    result = audit_docx(path)
    assert result["passed"] is False
    assert any("no_real_bullets" in v for v in result["violations"])


def test_excessive_bold_is_warning_not_violation(tmp_path):
    """A .docx with > 15 all-bold paragraphs → excessive_bold WARNING, not violation."""
    from api.cv.ats_audit import audit_docx

    path = _make_docx_all_bold(tmp_path)
    result = audit_docx(path)
    assert any("excessive_bold" in w for w in result["warnings"]), (
        f"Expected excessive_bold warning. warnings={result['warnings']}"
    )
    # excessive_bold is a WARNING — it should NOT appear in violations
    assert not any("excessive_bold" in v for v in result["violations"])


def test_no_context_lines_is_warning_not_violation(tmp_path):
    """A .docx without italic paragraphs → no_context_lines WARNING, not violation."""
    from api.cv.ats_audit import audit_docx

    path = _make_docx_no_bullets(tmp_path)  # also has no italics
    result = audit_docx(path)
    assert any("no_context_lines" in w for w in result["warnings"]), (
        f"Expected no_context_lines warning. warnings={result['warnings']}"
    )
    # no_context_lines is a WARNING — not in violations
    assert not any("no_context_lines" in v for v in result["violations"])


def test_no_tab_stops_is_warning_not_violation(tmp_path):
    """A .docx without tab stops → no_tab_stops WARNING, not violation."""
    from api.cv.ats_audit import audit_docx

    path = _make_docx_no_bullets(tmp_path)  # also has no tab stops
    result = audit_docx(path)
    assert any("no_tab_stops" in w for w in result["warnings"]), (
        f"Expected no_tab_stops warning. warnings={result['warnings']}"
    )
    # no_tab_stops is a WARNING — not in violations
    assert not any("no_tab_stops" in v for v in result["violations"])


def test_clean_docx_has_no_warnings(tmp_path):
    """A well-formatted Phase 6 .docx must have no warnings."""
    from api.cv.ats_audit import audit_docx

    path = _make_clean_docx(tmp_path)
    result = audit_docx(path)
    assert result["warnings"] == [], f"Expected no warnings for clean Phase 6 docx. Got: {result['warnings']}"


def test_stats_track_new_counters(tmp_path):
    """Stats must include bold_count, italic_count, tab_count for clean Phase 6 docx."""
    from api.cv.ats_audit import audit_docx

    path = _make_clean_docx(tmp_path)
    result = audit_docx(path)
    stats = result["stats"]
    # Phase 6 docx has context lines (italic) and tab stops
    assert stats["italic_count"] >= 1, "Expected at least 1 italic paragraph (context lines)"
    assert stats["tab_count"] >= 1, "Expected at least 1 tab stop (company/date lines)"
    assert stats["bold_count"] < 15, f"Expected < 15 all-bold paragraphs, got {stats['bold_count']}"
