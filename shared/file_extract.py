"""File text extraction for CV upload — PDF and DOCX.

SYNC RULE: _heading_prefix, _runs_to_markdown, _table_to_markdown, docx_to_markdown
are copied from api/onboard_utils.py. When changing either copy, update both.
"""

import io
import os
import tempfile
from pathlib import Path


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Extract plain text from a .pdf or .docx file.

    Args:
        file_bytes: Raw file content.
        filename: Original filename (used for extension detection).

    Returns:
        Extracted text as a string.

    Raises:
        ValueError: If the file extension is not .pdf or .docx.
    """
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(file_bytes)
    if ext == ".docx":
        return _extract_docx(file_bytes)
    raise ValueError(f"Unsupported file type: {ext!r}. Only .pdf and .docx are accepted.")


def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using pdfplumber."""
    import pdfplumber

    pages = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n".join(pages)


def _extract_docx(file_bytes: bytes) -> str:
    """Write bytes to a temp file and extract markdown text via docx_to_markdown."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name
    try:
        return docx_to_markdown(tmp_path)
    finally:
        os.unlink(tmp_path)


# ── DOCX helpers (synced copy from api/onboard_utils.py) ─────────────────


def _heading_prefix(para) -> str:
    """Return markdown heading prefix for a paragraph style, or ''."""
    style = para.style.name if para.style else ""
    if style == "Heading 1":
        return "# "
    if style == "Heading 2":
        return "## "
    if style == "Heading 3":
        return "### "
    if style == "Heading 4":
        return "#### "
    return ""


def _runs_to_markdown(para) -> str:
    """Convert paragraph runs to markdown, preserving bold."""
    parts = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        if run.bold:
            text = f"**{text}**"
        parts.append(text)
    return "".join(parts)


def _table_to_markdown(table) -> str:
    """Convert a docx table to a markdown table string."""
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n".join(rows)


def docx_to_markdown(path: str) -> str:
    """Read a .docx file and return a clean markdown representation."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")

    doc = Document(path)
    lines = []

    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            para = None
            for p in doc.paragraphs:
                if p._element is child:
                    para = p
                    break
            if para is None:
                continue

            text = _runs_to_markdown(para).strip()
            if not text:
                lines.append("")
                continue

            heading = _heading_prefix(para)
            style = para.style.name if para.style else ""

            if heading:
                lines.append(f"{heading}{text}")
            elif "List" in style:
                lines.append(f"- {text}")
            else:
                lines.append(text)

        elif tag == "tbl":
            for tbl in doc.tables:
                if tbl._element is child:
                    lines.append("")
                    lines.append(_table_to_markdown(tbl))
                    lines.append("")
                    break

    result = []
    blank_count = 0
    for line in lines:
        if line == "":
            blank_count += 1
            if blank_count <= 1:
                result.append(line)
        else:
            blank_count = 0
            result.append(line)

    return "\n".join(result).strip()
