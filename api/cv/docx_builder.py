"""ATS-compliant .docx builder — Phase 6 skill-quality formatting.

Parses structured markdown output from the LLM and builds a python-docx
Document that passes ATS parsers (Workday, Greenhouse, Lever) while
matching the visual quality of career-helper skill output.

Formatting spec:
  Name (#)            18pt bold  #1F4E79  after:60
  Title (line 2)      11pt       #444444  after:80
  Contact (line 3)     9pt       #555555  after:200
  Section (##)        11pt bold  #1F4E79  before:240 after:80
  Bullets (-)         10pt black          after:40   List Bullet style
  Context (_..._)     10pt italic #555555 after:60
  Company+Role+Date   10pt bold/muted     before:160 after:40  tab@9026
  Core Skills theme   10pt bold run + normal run
  Project name+URL    10pt bold + 9pt #555555
  Prose / education   10pt black          after:60/40

ATS constraints:
  - Zero tables anywhere in the document XML
  - Bullets via 'List Bullet' style (never unicode characters)
  - Single column, no tabs except right-aligned date
  - Post-processing: replaces em dash, en dash, arrows, Oxford comma
"""
import logging
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_TAB_ALIGNMENT

logger = logging.getLogger(__name__)

# ── Color palette ────────────────────────────────────────────────────────────
COLOR_ACCENT = RGBColor(0x1F, 0x4E, 0x79)  # dark blue: name, section headers
COLOR_TITLE  = RGBColor(0x44, 0x44, 0x44)  # dark gray: professional title
COLOR_MUTED  = RGBColor(0x55, 0x55, 0x55)  # medium gray: contact, dates, URLs, context
# Black = default (no color attribute): body text, bullets, company names

# Tab stop for right-aligned date: A4 content width at 1" margins
# A4 = 8.27", content = 8.27" - 2×1" = 6.27" ≈ 9026 twips
_DATE_TAB_POS = Inches(6.27)

# ATS post-processing: characters to replace
_REPLACEMENTS = [
    ("\u2014", "-"),   # em dash → hyphen
    ("\u2013", "-"),   # en dash → hyphen
    ("\u2192", ""),    # → arrow → removed
    ("=>",     ""),    # => arrow → removed
    ("\u007e", ""),    # tilde
    ("\u2022", ""),    # • bullet
    ("\u25e6", ""),    # ◦ hollow bullet
    ("\u25aa", ""),    # ▪ small square
]


def _clean_text(text: str) -> str:
    """Strip **bold** markers, apply ATS replacements, remove Oxford commas."""
    # Strip inline **bold** markers (builder handles bold via element type)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    # ATS character replacements
    for bad, good in _REPLACEMENTS:
        if bad in text:
            logger.warning("ATS post-processing: replaced %r in: %.60s", bad, text)
            text = text.replace(bad, good)
    # Oxford comma: ", and <word>" → ", <word>"
    cleaned = re.sub(r",\s+and\s+(?=\w)", ", ", text)
    if cleaned != text:
        logger.warning("ATS post-processing: removed Oxford comma in: %.60s", text)
        text = cleaned
    return text


def _strip_fences_and_preamble(markdown: str) -> str:
    """Remove code fences and any text before the first # heading."""
    markdown = re.sub(r"^```[a-zA-Z]*\n?", "", markdown.strip())
    markdown = re.sub(r"\n?```$", "", markdown.strip())
    match = re.search(r"^# ", markdown, re.MULTILINE)
    if not match:
        return markdown.strip()
    return markdown[match.start():].strip()


# ── Run / paragraph helpers ──────────────────────────────────────────────────

def _set_run(run, size_pt: int, bold: bool = False, italic: bool = False,
             color: RGBColor = None) -> None:
    """Configure font on a run (Calibri, size, bold, italic, color)."""
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if italic:
        run.font.italic = True
    if color is not None:
        run.font.color.rgb = color


def _set_spacing(para, before_twips: int = 0, after_twips: int = 0) -> None:
    """Set paragraph spacing in twips (1pt = 20 twips)."""
    if before_twips:
        para.paragraph_format.space_before = Pt(before_twips / 20)
    if after_twips:
        para.paragraph_format.space_after = Pt(after_twips / 20)


# ── Element builders ─────────────────────────────────────────────────────────

def _add_name(doc: Document, text: str) -> None:
    """Name: 18pt bold #1F4E79, after:60."""
    para = doc.add_paragraph()
    run = para.add_run(_clean_text(text))
    _set_run(run, 18, bold=True, color=COLOR_ACCENT)
    _set_spacing(para, after_twips=60)


def _add_title_line(doc: Document, text: str) -> None:
    """Professional title: 11pt #444444, after:80."""
    para = doc.add_paragraph()
    run = para.add_run(_clean_text(text))
    _set_run(run, 11, color=COLOR_TITLE)
    _set_spacing(para, after_twips=80)


def _add_contact_line(doc: Document, text: str) -> None:
    """Contact info: 9pt #555555, after:200."""
    para = doc.add_paragraph()
    run = para.add_run(_clean_text(text))
    _set_run(run, 9, color=COLOR_MUTED)
    _set_spacing(para, after_twips=200)


def _add_section_header(doc: Document, text: str) -> None:
    """Section header (##): 11pt bold #1F4E79, before:240, after:80."""
    para = doc.add_paragraph()
    run = para.add_run(_clean_text(text))
    _set_run(run, 11, bold=True, color=COLOR_ACCENT)
    _set_spacing(para, before_twips=240, after_twips=80)


def _add_bullet(doc: Document, text: str) -> None:
    """Bullet (- ): 10pt black, List Bullet style, after:40."""
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(_clean_text(text))
    _set_run(run, 10)
    _set_spacing(para, after_twips=40)


def _add_italic_context(doc: Document, text: str) -> None:
    """Context line (_text_): 10pt italic #555555, after:60."""
    para = doc.add_paragraph()
    run = para.add_run(_clean_text(text))
    _set_run(run, 10, italic=True, color=COLOR_MUTED)
    _set_spacing(para, after_twips=60)


def _add_company_date_line(doc: Document, company_role: str, date: str) -> None:
    """Company+Role (bold black) + tab + Date (#555555). Tab stop at 9026 twips.

    Two runs: run1 = company/role (10pt bold black), run2 = \\tdate (10pt #555555).
    before:160, after:40.
    """
    para = doc.add_paragraph()
    para.paragraph_format.tab_stops.add_tab_stop(_DATE_TAB_POS, WD_TAB_ALIGNMENT.RIGHT)
    _set_spacing(para, before_twips=160, after_twips=40)
    run1 = para.add_run(_clean_text(company_role))
    _set_run(run1, 10, bold=True)
    run2 = para.add_run("\t" + _clean_text(date))
    _set_run(run2, 10, color=COLOR_MUTED)


def _add_core_skills_line(doc: Document, theme: str, prose: str) -> None:
    """Core Skills: bold theme run + normal prose run. 10pt black, after:80."""
    para = doc.add_paragraph()
    run_theme = para.add_run(_clean_text(theme) + ": ")
    _set_run(run_theme, 10, bold=True)
    if prose:
        run_prose = para.add_run(_clean_text(prose))
        _set_run(run_prose, 10)
    _set_spacing(para, after_twips=80)


def _add_project_line(doc: Document, name: str, url: str = "") -> None:
    """Project: name (10pt bold black) + optional URL (9pt #555555). after:40."""
    para = doc.add_paragraph()
    run_name = para.add_run(_clean_text(name))
    _set_run(run_name, 10, bold=True)
    if url:
        run_url = para.add_run("  " + url.strip())
        _set_run(run_url, 9, color=COLOR_MUTED)
    _set_spacing(para, after_twips=40)


def _add_project_description(doc: Document, text: str) -> None:
    """Project description: 10pt black, after:100."""
    para = doc.add_paragraph()
    run = para.add_run(_clean_text(text))
    _set_run(run, 10)
    _set_spacing(para, after_twips=100)


def _add_prose(doc: Document, text: str) -> None:
    """General prose paragraph: 10pt black, after:60."""
    para = doc.add_paragraph()
    run = para.add_run(_clean_text(text))
    _set_run(run, 10)
    _set_spacing(para, after_twips=60)


def _add_education_line(doc: Document, text: str) -> None:
    """Education / language line: 10pt black, after:40."""
    para = doc.add_paragraph()
    run = para.add_run(_clean_text(text))
    _set_run(run, 10)
    _set_spacing(para, after_twips=40)


# ── Section parsers ──────────────────────────────────────────────────────────

def _is_italic_context_line(line: str) -> bool:
    """Return True if line is a standalone _italic context_ marker."""
    s = line.strip()
    return (
        s.startswith("_")
        and s.endswith("_")
        and len(s) > 2
        and not s.startswith("__")
    )


def _parse_work_experience(doc: Document, lines: list[str]) -> None:
    """Parse Work Experience section.

    Handles Phase 6 format:
        ### Company - Role\\tDate
        _context line_
        - bullet

    Also handles Phase 5 backward compat:
        ### Company Name
        **Role | Date**
        - bullet
    """
    for line in lines:
        s = line.rstrip()
        if s.startswith("### "):
            content = s[4:]
            if "\t" in content:
                # Phase 6: tab separates role from date
                parts = content.split("\t", 1)
                _add_company_date_line(doc, parts[0].strip(), parts[1].strip())
            else:
                # Phase 5 compat: company name only (date on next line)
                _add_company_date_line(doc, content.strip(), "")
        elif _is_italic_context_line(s):
            _add_italic_context(doc, s[1:-1])
        elif s.startswith("**") and s.endswith("**") and "|" in s:
            # Phase 5 compat: **Role | Date** → render as company line
            inner = s.strip("*").strip()
            parts = inner.split("|", 1)
            _add_company_date_line(doc, parts[0].strip(), parts[1].strip() if len(parts) > 1 else "")
        elif s.startswith("- "):
            _add_bullet(doc, s[2:])
        elif s.strip():
            _add_prose(doc, s)


def _parse_projects(doc: Document, lines: list[str]) -> None:
    """Parse Projects section.

    Phase 6 format:
        ### ProjectName  URL
        description line(s)
        - bullet (optional)

    Phase 5 compat: ### ProjectName / description line
    """
    i = 0
    in_project = False
    while i < len(lines):
        s = lines[i].rstrip()
        if s.startswith("### "):
            content = s[4:]
            # URL separated by 2+ spaces
            parts = re.split(r"  +", content, maxsplit=1)
            if len(parts) == 2:
                _add_project_line(doc, parts[0].strip(), parts[1].strip())
            else:
                _add_project_line(doc, content.strip())
            in_project = True
        elif s.startswith("- "):
            _add_bullet(doc, s[2:])
        elif s.strip():
            if in_project:
                _add_project_description(doc, s)
            else:
                _add_prose(doc, s)
        i += 1


def _parse_core_skills(doc: Document, lines: list[str]) -> None:
    """Parse Core Skills section.

    Phase 6 format: "Theme Name: prose text on same line"
    Phase 5 compat: **Theme Name** / prose paragraph on next line(s)
    """
    i = 0
    while i < len(lines):
        s = lines[i].rstrip()
        if not s.strip():
            i += 1
            continue

        # Phase 6: "Theme: prose" — colon within first ~60 chars, not a bullet/bold
        colon_match = re.match(r"^([A-Za-z][^:*\n]{0,50}):\s+(.+)", s)
        if colon_match and not s.startswith("- ") and not s.startswith("**"):
            _add_core_skills_line(doc, colon_match.group(1).strip(),
                                  colon_match.group(2).strip())
        elif s.startswith("**") and s.endswith("**"):
            # Phase 5 compat: **Theme Name** block
            theme = s.strip("*").strip()
            prose_lines = []
            j = i + 1
            while j < len(lines) and lines[j].strip() and not lines[j].startswith("**"):
                prose_lines.append(lines[j].strip())
                j += 1
            _add_core_skills_line(doc, theme, " ".join(prose_lines))
            i = j
            continue
        elif s.startswith("- "):
            _add_bullet(doc, s[2:])
        elif s.strip():
            _add_prose(doc, s)
        i += 1


def _parse_bullets_section(doc: Document, lines: list[str]) -> None:
    """Parse a section that is primarily bullets (Selected Impact, Education, Languages)."""
    for line in lines:
        s = line.rstrip()
        if s.startswith("- "):
            _add_bullet(doc, s[2:])
        elif s.strip():
            _add_education_line(doc, s)


def _parse_prose_section(doc: Document, lines: list[str]) -> None:
    """Parse Summary and other prose sections."""
    for line in lines:
        s = line.rstrip()
        if s.startswith("- "):
            _add_bullet(doc, s[2:])
        elif s.strip():
            _add_prose(doc, s)


# ── Main entry point ─────────────────────────────────────────────────────────

def build_docx(markdown: str, output_path: str) -> str:
    """Parse structured LLM markdown and build an ATS-compliant .docx file.

    Supports both Phase 6 (new contract with tab-dates and _context lines_)
    and Phase 5 (old contract with **Role | Date** lines) for backward compat.

    Args:
        markdown: Structured CV markdown string from the LLM.
        output_path: Absolute path where the .docx file will be written.

    Returns:
        The output_path on success.

    Raises:
        ValueError: If markdown is empty or has no # heading.
    """
    if not markdown or not markdown.strip():
        raise ValueError("Cannot build .docx from empty markdown.")

    markdown = _strip_fences_and_preamble(markdown)

    if not re.search(r"^# ", markdown, re.MULTILINE):
        raise ValueError(
            "Malformed CV markdown: no '# Name' heading found. "
            "The LLM output must start with a level-1 heading."
        )

    doc = Document()

    # Page setup: A4, 1-inch margins
    section = doc.sections[0]
    section.page_width  = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Default paragraph style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    lines = markdown.split("\n")

    # ── Header block (before first ## section) ────────────────────────────
    header_lines = []
    section_start = 0
    for idx, line in enumerate(lines):
        if line.startswith("## "):
            section_start = idx
            break
        header_lines.append(line)

    # First non-empty → name, second → title, subsequent → contact
    non_empty_header = [l for l in header_lines if l.strip() and not l.startswith("#")]
    name_line = next(
        (l.lstrip("# ").strip() for l in header_lines if l.startswith("# ")),
        non_empty_header[0].strip() if non_empty_header else "",
    )

    if name_line:
        _add_name(doc, name_line)

    # Title and contact: lines after the # Name line, before first ##
    extra = [l for l in header_lines if l.strip() and not l.startswith("#")]
    for idx, line in enumerate(extra):
        if idx == 0:
            _add_title_line(doc, line.strip())
        else:
            _add_contact_line(doc, line.strip())

    # ── Sections ──────────────────────────────────────────────────────────
    sections: list[tuple[str, list[str]]] = []
    current_name = ""
    current_lines: list[str] = []

    for line in lines[section_start:]:
        if line.startswith("## "):
            if current_name:
                sections.append((current_name, current_lines))
            current_name = line[3:].strip()
            current_lines = []
        else:
            current_lines.append(line)
    if current_name:
        sections.append((current_name, current_lines))

    for sec_name, sec_lines in sections:
        _add_section_header(doc, sec_name)
        # Strip leading/trailing blank lines from section body
        while sec_lines and not sec_lines[0].strip():
            sec_lines.pop(0)
        while sec_lines and not sec_lines[-1].strip():
            sec_lines.pop()

        sec_lower = sec_name.lower()
        if "work experience" in sec_lower:
            _parse_work_experience(doc, sec_lines)
        elif "project" in sec_lower:
            _parse_projects(doc, sec_lines)
        elif "core skills" in sec_lower or "skills" in sec_lower:
            _parse_core_skills(doc, sec_lines)
        elif any(k in sec_lower for k in ("impact", "education", "language", "certification")):
            _parse_bullets_section(doc, sec_lines)
        else:
            _parse_prose_section(doc, sec_lines)

    doc.save(output_path)
    return output_path
