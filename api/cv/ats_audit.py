"""ATS compliance audit utility.

Post-build safety net that inspects a .docx file for ATS violations.
Called after docx_builder.build_docx() — result exposed via X-ATS-Audit header.

Phase 6 additions:
  - Counts all-bold paragraphs: > 15 → warning "excessive_bold"
  - Checks for at least 1 italic paragraph → missing = warning "no_context_lines"
  - Checks for bullet numbering (List style) → missing = VIOLATION "no_real_bullets"
  - Checks for at least 1 tab stop → missing = warning "no_tab_stops"
"""

import re
from docx import Document
from docx.oxml.ns import qn

# Standard section headers required by the CV format
_REQUIRED_HEADERS = {
    "Summary",
    "Selected Impact",
    "Core Skills",
    "Work Experience",
    "Education and Certifications",
    "Languages",
}

# Prohibited characters with their violation labels
_PROHIBITED_CHARS = [
    ("\u2014", "prohibited_char:em_dash"),
    ("\u2013", "prohibited_char:en_dash"),
    ("\u2192", "prohibited_char:arrow"),
    ("\u25cf", "prohibited_char:unicode_bullet"),
    ("\u25e6", "prohibited_char:unicode_bullet"),
    ("\u25aa", "prohibited_char:unicode_bullet"),
    ("\u2022", "prohibited_char:unicode_bullet"),
]

# Approximate lines per page for page count heuristic
_LINES_PER_PAGE = 45

# Threshold for "excessive bold" warning
_BOLD_PARA_LIMIT = 15


def _para_is_all_bold(para) -> bool:
    """Return True if every non-empty run in the paragraph is explicitly bold (via XML)."""
    text_runs = [r for r in para.runs if r.text.strip()]
    if not text_runs:
        return False
    for run in text_runs:
        rPr = run._r.find(qn("w:rPr"))
        if rPr is None:
            return False
        b_el = rPr.find(qn("w:b"))
        if b_el is None:
            return False
        # <w:b w:val="0"/> means explicitly NOT bold
        val = b_el.get(qn("w:val"))
        if val is not None and val.lower() in ("0", "false", "off"):
            return False
    return True


def _para_has_italic(para) -> bool:
    """Return True if any run in the paragraph has explicit italic via XML."""
    for run in para.runs:
        rPr = run._r.find(qn("w:rPr"))
        if rPr is not None and rPr.find(qn("w:i")) is not None:
            # Check it's not explicitly disabled
            i_el = rPr.find(qn("w:i"))
            val = i_el.get(qn("w:val"))
            if val is None or val.lower() not in ("0", "false", "off"):
                return True
    return False


def _para_has_tab_stop(para) -> bool:
    """Return True if the paragraph has a <w:tabs> element in its pPr."""
    pPr = para._p.pPr
    return pPr is not None and pPr.find(qn("w:tabs")) is not None


def audit_docx(path: str) -> dict:
    """Inspect a .docx file for ATS compliance violations and formatting quality warnings.

    Args:
        path: Absolute path to the .docx file to audit.

    Returns:
        Dict with keys:
          - passed (bool): True if zero violations found (warnings don't affect passed).
          - violations (list[str]): ATS violation codes/descriptions.
          - warnings (list[str]): Formatting quality warnings (don't fail the audit).
          - stats (dict): section_count, bullet_count, paragraph_count,
                         estimated_pages, bold_count, italic_count, tab_count.
    """
    violations: list[str] = []
    warnings: list[str] = []
    doc = Document(path)

    # 1. Zero tables
    if len(doc.tables) > 0:
        violations.append(f"table_found: {len(doc.tables)} table(s) in document")

    # 2. Collect all paragraph text and inspect
    all_paragraphs = doc.paragraphs
    paragraph_count = len(all_paragraphs)

    section_count = 0
    bullet_count = 0
    italic_count = 0
    tab_count = 0
    bold_count = 0
    found_headers: set[str] = set()

    for para in all_paragraphs:
        text = para.text
        stripped = text.strip()

        # Detect section headers: "Heading 2" style (Phase 5) OR exact text match
        # against required headers (Phase 6 uses Normal style with explicit formatting)
        is_section_header = para.style.name == "Heading 2" or stripped in _REQUIRED_HEADERS
        if is_section_header:
            section_count += 1
            for required in _REQUIRED_HEADERS:
                if required.lower() in stripped.lower():
                    found_headers.add(required)

        # Count bullets (List Bullet style = real numbering, not unicode)
        if "List" in para.style.name or "Bullet" in para.style.name:
            bullet_count += 1

        # Count all-bold paragraphs
        if _para_is_all_bold(para):
            bold_count += 1

        # Count italic paragraphs (context lines)
        if _para_has_italic(para):
            italic_count += 1

        # Count paragraphs with tab stops (date lines)
        if _para_has_tab_stop(para):
            tab_count += 1

        # Check for prohibited characters
        for char, label in _PROHIBITED_CHARS:
            if char in text:
                violations.append(f"{label}: found in paragraph: {text[:60]}")
                break  # one violation per paragraph

        # Oxford comma check
        if re.search(r",\s+and\s+\w", text):
            violations.append(f"oxford_comma: found in paragraph: {text[:60]}")

    # 3. Check date format in role lines (Phase 5: | separator; Phase 6: \t separator)
    for para in all_paragraphs:
        if para.style.name in ("Normal", "Body Text") and "|" in para.text:
            date_part = para.text.split("|", 1)[-1].strip()
            if date_part:
                if not re.match(
                    r"^(\d{2}/\d{4}|\d{4})\s*[-\u2013]\s*(\d{2}/\d{4}|\d{4}|[Pp]resent)$",
                    date_part.strip(),
                ):
                    if re.search(r"\d{4}", date_part):
                        violations.append(f"date_format: unexpected date format: {date_part[:40]}")

    # 4. Missing required headers
    for required in _REQUIRED_HEADERS:
        if required not in found_headers:
            violations.append(f"missing_header:{required}")

    # 5. Phase 6 formatting quality checks

    # no_real_bullets → VIOLATION (ATS parsers need real bullets)
    if bullet_count == 0:
        violations.append("no_real_bullets: no List Bullet paragraphs found")

    # excessive_bold → WARNING only
    if bold_count > _BOLD_PARA_LIMIT:
        warnings.append(
            f"excessive_bold: {bold_count} all-bold paragraphs "
            f"(expected < {_BOLD_PARA_LIMIT}, only headers/company lines should be bold)"
        )

    # no_context_lines → WARNING only
    if italic_count == 0:
        warnings.append(
            "no_context_lines: no italic paragraphs found (expected _role context lines_ under each company)"
        )

    # no_tab_stops → WARNING only
    if tab_count == 0:
        warnings.append(
            "no_tab_stops: no paragraphs with tab stops found (expected right-aligned dates on company lines)"
        )

    # 6. Stats
    estimated_pages = max(1, round(paragraph_count / _LINES_PER_PAGE + 0.4))
    stats = {
        "section_count": section_count,
        "bullet_count": bullet_count,
        "paragraph_count": paragraph_count,
        "estimated_pages": estimated_pages,
        "bold_count": bold_count,
        "italic_count": italic_count,
        "tab_count": tab_count,
    }

    return {
        "passed": len(violations) == 0,
        "violations": violations,
        "warnings": warnings,
        "stats": stats,
    }
