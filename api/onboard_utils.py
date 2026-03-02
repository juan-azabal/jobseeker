"""Onboarding utilities extracted from agent/onboard.py for API use.

Provides CV-to-markdown conversion, LLM profile extraction, and profile
YAML generation without any sys.path dependency on the agent directory.

Note: This is the API-side copy. The agent has its own version at
agent/onboard.py. If you change shared logic, update both.
"""

import json
from pathlib import Path

import yaml

from api import config

# ---------------------------------------------------------------------------
# Extraction prompt (loaded once, cached)
# ---------------------------------------------------------------------------

_EXTRACTION_PROMPT: str | None = None
_PROMPT_PATH = Path(__file__).parent / "prompts" / "onboard-extraction.md"


def _load_extraction_prompt() -> str:
    global _EXTRACTION_PROMPT
    if _EXTRACTION_PROMPT is None:
        with open(_PROMPT_PATH) as f:
            _EXTRACTION_PROMPT = f.read()
    return _EXTRACTION_PROMPT


# ---------------------------------------------------------------------------
# .docx -> markdown
# ---------------------------------------------------------------------------


def _heading_prefix(para) -> str:
    """Return markdown heading prefix for a paragraph style, or ''."""
    style = para.style.name if para.style else ""
    if style == "Heading 1":
        return "# "
    if style == "Heading 2":
        return "## "
    if style == "Heading 3":
        return "### "
    if style == "Heading 4":
        return "#### "
    return ""


def _runs_to_markdown(para) -> str:
    """Convert paragraph runs to markdown, preserving bold."""
    parts = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        if run.bold:
            text = f"**{text}**"
        parts.append(text)
    return "".join(parts)


def _table_to_markdown(table) -> str:
    """Convert a docx table to a markdown table string."""
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n".join(rows)


def docx_to_markdown(path: str) -> str:
    """Read a .docx file and return a clean markdown representation."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")

    doc = Document(path)
    lines = []

    # Iterate over block-level elements in document order (paragraphs + tables)
    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            para = None
            for p in doc.paragraphs:
                if p._element is child:
                    para = p
                    break
            if para is None:
                continue

            text = _runs_to_markdown(para).strip()
            if not text:
                lines.append("")
                continue

            heading = _heading_prefix(para)
            style = para.style.name if para.style else ""

            if heading:
                lines.append(f"{heading}{text}")
            elif "List" in style:
                lines.append(f"- {text}")
            else:
                lines.append(text)

        elif tag == "tbl":
            for tbl in doc.tables:
                if tbl._element is child:
                    lines.append("")
                    lines.append(_table_to_markdown(tbl))
                    lines.append("")
                    break

    # Collapse consecutive blank lines to at most one
    result = []
    blank_count = 0
    for line in lines:
        if line == "":
            blank_count += 1
            if blank_count <= 1:
                result.append(line)
        else:
            blank_count = 0
            result.append(line)

    return "\n".join(result).strip()


# ---------------------------------------------------------------------------
# LLM extraction
# ---------------------------------------------------------------------------


def _extract_profile(cv_text: str, client) -> dict:
    """Call gpt-4o-mini to extract profile fields from CV text. Retries once.

    Args:
        cv_text: Markdown text of the CV.
        client: An OpenAI client instance.

    Raises:
        RuntimeError: If extraction fails after retries.
    """
    prompt = _load_extraction_prompt()
    messages = [
        {"role": "system", "content": prompt},
        {"role": "user", "content": cv_text[:12000]},  # cap at ~3K tokens
    ]

    for attempt in range(2):
        try:
            response = client.chat.completions.create(
                model=config.LLM_MODEL_PARSING,
                messages=messages,
                temperature=0,
                max_tokens=1500,
            )
            raw = response.choices[0].message.content.strip()

            # Strip markdown fences if model added them
            if raw.startswith("```"):
                raw = raw.split("\n", 1)[1] if "\n" in raw else raw[3:]
            if raw.endswith("```"):
                raw = raw[:-3]
            raw = raw.strip()

            return json.loads(raw)

        except json.JSONDecodeError as e:
            if attempt == 0:
                messages[0]["content"] += (
                    "\n\nCRITICAL: Return ONLY the JSON object. No text before or after. No markdown fences."
                )
            else:
                raise RuntimeError(f"Could not parse LLM response as JSON: {e}")

        except Exception as e:
            if attempt == 0:
                continue
            else:
                raise RuntimeError(f"LLM extraction failed: {e}")


# ---------------------------------------------------------------------------
# Profile YAML generation
# ---------------------------------------------------------------------------


def _build_profile_yaml(
    extracted: dict,
    profile_id: str,
    email: str,
    salary_min: int,
    location_choice: str,
    home_locations: list[str],
) -> str:
    """Build the profile YAML string from extracted data + user answers."""

    profile: dict = {
        "user": {
            "id": profile_id,
            "name": extracted["name"],
            "email": email,
            "active": True,
            "languages": extracted.get("languages", ["en"]),
            "home_locations": home_locations,
            "location_preference": location_choice,
        },
        "target": {
            "level": extracted["target_level"],
            "track": extracted["track"],
            **({"role_type": extracted["role_type"]} if extracted.get("role_type") else {}),
            **({"role_function": extracted["role_function"]} if extracted.get("role_function") else {}),
            "domains": extracted.get("domains", {}),
        },
        "scoring": {
            "rag_threshold": 25,
            "salary_min": salary_min,
        },
        "skills": extracted.get("skills", []),
        "knowledge": {
            "dir": f"knowledge/{profile_id}",
            "profile_files": ["cv.md"],
            "collection_name": f"{profile_id}_profile",
        },
        "stories": {},
        "preferences": "config/preferences.yaml",
        "searches": "config/searches.yaml",
        "watchlist": "config/watchlist.yaml",
        "seen_ids": f"config/seen_ids/{profile_id}.txt",
    }

    exclude = extracted.get("exclude_companies", [])
    if exclude:
        profile["exclude_companies"] = exclude

    header = f"# JobAgent profile \u2014 {extracted['name']}\n# Generated by onboard.py\n\n"
    return header + yaml.dump(profile, default_flow_style=False, allow_unicode=True, sort_keys=False)
