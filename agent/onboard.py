"""
CV onboarding: .docx → profile YAML + knowledge/cv.md in under 5 minutes.

Usage:
    python onboard.py --cv path/to/cv.docx
    python onboard.py --cv path/to/cv.docx --profile custom_id

Output:
    config/profiles/{id}.yaml
    knowledge/{id}/cv.md
    config/seen_ids/{id}.txt
"""

import argparse
import json
import os
import sys

import yaml
from dotenv import load_dotenv
from openai import OpenAI

load_dotenv()

# ---------------------------------------------------------------------------
# Extraction prompt (loaded once, cached)
# ---------------------------------------------------------------------------

_EXTRACTION_PROMPT: str | None = None
_PROMPT_PATH = os.path.join(os.path.dirname(__file__), "prompts", "onboard-extraction.md")


def _load_extraction_prompt() -> str:
    global _EXTRACTION_PROMPT
    if _EXTRACTION_PROMPT is None:
        with open(_PROMPT_PATH) as f:
            _EXTRACTION_PROMPT = f.read()
    return _EXTRACTION_PROMPT


# ---------------------------------------------------------------------------
# .docx → markdown
# ---------------------------------------------------------------------------


def _heading_prefix(para) -> str:
    """Return markdown heading prefix for a paragraph style, or ''."""
    style = para.style.name if para.style else ""
    if style == "Heading 1":
        return "# "
    if style == "Heading 2":
        return "## "
    if style == "Heading 3":
        return "### "
    if style == "Heading 4":
        return "#### "
    return ""


def _runs_to_markdown(para) -> str:
    """Convert paragraph runs to markdown, preserving bold."""
    parts = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        if run.bold:
            text = f"**{text}**"
        parts.append(text)
    return "".join(parts)


def _table_to_markdown(table) -> str:
    """Convert a docx table to a markdown table string."""
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n".join(rows)


def docx_to_markdown(path: str) -> str:
    """Read a .docx file and return a clean markdown representation."""
    try:
        from docx import Document
    except ImportError:
        print("Error: python-docx is not installed. Run: pip install python-docx")
        sys.exit(1)

    try:
        doc = Document(path)
    except Exception as e:
        print(f"Error: Could not read {path}: {e}")
        print("Check that the file is a valid .docx (not corrupted or password-protected).")
        sys.exit(1)

    lines = []

    # Iterate over block-level elements in document order (paragraphs + tables)
    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            para = None
            for p in doc.paragraphs:
                if p._element is child:
                    para = p
                    break
            if para is None:
                continue

            text = _runs_to_markdown(para).strip()
            if not text:
                lines.append("")
                continue

            heading = _heading_prefix(para)
            style = para.style.name if para.style else ""

            if heading:
                lines.append(f"{heading}{text}")
            elif "List" in style:
                lines.append(f"- {text}")
            else:
                lines.append(text)

        elif tag == "tbl":
            for tbl in doc.tables:
                if tbl._element is child:
                    lines.append("")
                    lines.append(_table_to_markdown(tbl))
                    lines.append("")
                    break

    # Collapse consecutive blank lines to at most one
    result = []
    blank_count = 0
    for line in lines:
        if line == "":
            blank_count += 1
            if blank_count <= 1:
                result.append(line)
        else:
            blank_count = 0
            result.append(line)

    return "\n".join(result).strip()


# ---------------------------------------------------------------------------
# LLM extraction
# ---------------------------------------------------------------------------


def _extract_profile(cv_text: str, client: OpenAI) -> dict:
    """Call gpt-4o-mini to extract profile fields from CV text. Retries once."""
    prompt = _load_extraction_prompt()
    messages = [
        {"role": "system", "content": prompt},
        {"role": "user", "content": cv_text[:12000]},  # cap at ~3K tokens
    ]

    for attempt in range(2):
        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=messages,
                temperature=0,
                max_tokens=1500,
            )
            raw = response.choices[0].message.content.strip()

            # Strip markdown fences if model added them
            if raw.startswith("```"):
                raw = raw.split("\n", 1)[1] if "\n" in raw else raw[3:]
            if raw.endswith("```"):
                raw = raw[:-3]
            raw = raw.strip()

            return json.loads(raw)

        except json.JSONDecodeError as e:
            if attempt == 0:
                print("   Warning: JSON parse error on first attempt, retrying...")
                messages[0]["content"] += (
                    "\n\nCRITICAL: Return ONLY the JSON object. No text before or after. No markdown fences."
                )
            else:
                print(f"   Error: Could not parse LLM response as JSON: {e}")
                print(f"   Raw response:\n{raw}")
                sys.exit(1)

        except Exception as e:
            if attempt == 0:
                print(f"   Warning: API call failed ({e}), retrying...")
            else:
                print(f"   Error: LLM extraction failed: {e}")
                print("   Check that OPENAI_API_KEY is set in .env")
                sys.exit(1)


# ---------------------------------------------------------------------------
# Interactive questions
# ---------------------------------------------------------------------------


def _ask_salary() -> int:
    """Ask for minimum salary in EUR. Validates it's a reasonable integer."""
    while True:
        raw = input("1. Minimum salary (EUR, e.g. 80000): ").strip()
        try:
            val = int(raw.replace(",", "").replace(".", ""))
            if 10000 <= val <= 1000000:
                return val
            print("   Please enter a value between 10,000 and 1,000,000.")
        except ValueError:
            print("   Please enter a number (e.g. 80000).")


_COUNTRY_CITIES = {
    "spain": ["madrid", "barcelona", "valencia", "seville", "bilbao", "malaga"],
    "españa": ["madrid", "barcelona", "valencia", "seville", "bilbao", "malaga"],
    "france": ["paris", "lyon", "marseille", "bordeaux", "toulouse", "nantes"],
    "germany": ["berlin", "munich", "hamburg", "frankfurt", "cologne", "stuttgart"],
    "netherlands": ["amsterdam", "rotterdam", "utrecht", "eindhoven", "den haag"],
    "poland": ["warsaw", "krakow", "wroclaw", "gdansk", "poznan", "lodz"],
    "portugal": ["lisbon", "porto", "braga", "faro"],
    "italy": ["milan", "rome", "turin", "florence", "bologna", "naples"],
    "uk": ["london", "manchester", "edinburgh", "bristol", "cambridge"],
}


def _ask_location(home_locations: list[str]) -> tuple[str, list[str]]:
    """Ask location preference. Returns (choice, updated home_locations)."""
    city = home_locations[0].title() if home_locations else "your city"
    country = home_locations[1].title() if len(home_locations) > 1 else "your country"

    print("2. Location preference:")
    print("   a) Remote only")
    print(f"   b) Remote + {city} (hybrid/onsite in {city})")
    print(f"   c) Anywhere in {country}")
    print("   d) Anywhere in Europe")

    while True:
        raw = input("   Choice [a/b/c/d]: ").strip().lower()
        if raw in ("a", "b", "c", "d"):
            break
        print("   Please enter a, b, c, or d.")

    updated = list(home_locations)

    if raw == "c":
        for loc in home_locations:
            if loc in _COUNTRY_CITIES:
                for c in _COUNTRY_CITIES[loc]:
                    if c not in updated:
                        updated.append(c)

    # "d" (Europe): keep home_locations as-is; location_preference field signals broad accept

    return raw, updated


def _ask_email(default: str | None = None) -> str:
    """Ask for notification email. Uses CV email as default if available."""
    prompt = f"3. Email for daily digest [{default}]: " if default else "3. Email for daily digest: "
    while True:
        raw = input(prompt).strip()
        if not raw and default:
            return default
        if "@" in raw and "." in raw.split("@")[-1]:
            return raw
        print("   Please enter a valid email address.")


# ---------------------------------------------------------------------------
# Profile ID helpers
# ---------------------------------------------------------------------------


def _generate_profile_id(name: str) -> str:
    """Derive profile ID from first name, lowercase, ASCII-safe."""
    first = name.strip().split()[0].lower()
    safe = "".join(c for c in first if c.isalnum())
    return safe or "user"


# ---------------------------------------------------------------------------
# Profile YAML generation
# ---------------------------------------------------------------------------


def _build_search_titles(extracted: dict) -> list[str]:
    """Generate initial search_titles from role_type + level."""
    role_type = (extracted.get("role_type") or "").strip()
    level = (extracted.get("target_level") or "").strip()
    if not role_type:
        return []
    titles = []
    if level:
        titles.append(f"{level.capitalize()} {role_type}")
    titles.append(role_type)
    return titles


def _build_profile_yaml(
    extracted: dict,
    profile_id: str,
    email: str,
    salary_min: int,
    location_choice: str,
    home_locations: list[str],
) -> str:
    """Build the profile YAML string from extracted data + user answers."""

    profile: dict = {
        "user": {
            "id": profile_id,
            "name": extracted["name"],
            "email": email,
            "active": True,
            "languages": extracted.get("languages", ["en"]),
            "home_locations": home_locations,
            "location_preference": location_choice,
        },
        "target": {
            "level": extracted["target_level"],
            "track": extracted["track"],
            **({"role_type": extracted["role_type"]} if extracted.get("role_type") else {}),
            **({"role_function": extracted["role_function"]} if extracted.get("role_function") else {}),
            "search_titles": _build_search_titles(extracted),
            "domains": extracted.get("domains", {}),
        },
        "scoring": {
            "rag_threshold": 25,
            "salary_min": salary_min,
        },
        "skills": extracted.get("skills", []),
        "knowledge": {
            "dir": f"knowledge/{profile_id}",
            "profile_files": ["cv.md"],
            "collection_name": f"{profile_id}_profile",
        },
        "stories": {},
        "preferences": "config/preferences.yaml",
        "watchlist": "config/watchlist.yaml",
        "seen_ids": f"config/seen_ids/{profile_id}.txt",
    }

    exclude = extracted.get("exclude_companies", [])
    if exclude:
        profile["exclude_companies"] = exclude

    header = f"# JobAgent profile — {extracted['name']}\n# Generated by onboard.py\n\n"
    return header + yaml.dump(profile, default_flow_style=False, allow_unicode=True, sort_keys=False)


# ---------------------------------------------------------------------------
# File I/O
# ---------------------------------------------------------------------------


def _save_files(profile_id: str, cv_markdown: str, profile_yaml: str) -> tuple[str, str, str]:
    """Write all generated files. Returns (profile_path, cv_path, seen_ids_path)."""

    profiles_dir = "config/profiles"
    os.makedirs(profiles_dir, exist_ok=True)
    profile_path = os.path.join(profiles_dir, f"{profile_id}.yaml")

    knowledge_dir = os.path.join("knowledge", profile_id)
    os.makedirs(knowledge_dir, exist_ok=True)
    cv_path = os.path.join(knowledge_dir, "cv.md")

    seen_ids_dir = "config/seen_ids"
    os.makedirs(seen_ids_dir, exist_ok=True)
    seen_ids_path = os.path.join(seen_ids_dir, f"{profile_id}.txt")

    with open(profile_path, "w") as f:
        f.write(profile_yaml)

    with open(cv_path, "w") as f:
        f.write(cv_markdown)

    if not os.path.exists(seen_ids_path):
        open(seen_ids_path, "w").close()

    return profile_path, cv_path, seen_ids_path


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def main():
    parser = argparse.ArgumentParser(description="Onboard a new user from a CV (.docx)")
    parser.add_argument("--cv", required=True, help="Path to the CV file (.docx)")
    parser.add_argument("--profile", default=None, help="Override the auto-generated profile ID")
    args = parser.parse_args()

    cv_path = args.cv

    if not os.path.exists(cv_path):
        print(f"Error: File not found: {cv_path}")
        sys.exit(1)

    if not cv_path.lower().endswith(".docx"):
        print("Error: Only .docx files are supported.")
        print("Export your CV from Google Docs or Word as .docx.")
        sys.exit(1)

    client = OpenAI()

    # Step 1: Extract text from .docx
    filename = os.path.basename(cv_path)
    print(f"Extracting text from {filename}...")
    cv_markdown = docx_to_markdown(cv_path)

    if len(cv_markdown.strip()) < 200:
        print("Error: CV text is too short. Check that the file has readable content.")
        sys.exit(1)

    # Step 2: LLM extraction
    print("Extracting profile from CV...")
    extracted = _extract_profile(cv_markdown, client)

    # Determine profile ID
    profile_id = args.profile or _generate_profile_id(extracted.get("name", "user"))

    # Check for existing profile
    profile_path = os.path.join("config/profiles", f"{profile_id}.yaml")
    if os.path.exists(profile_path):
        answer = input(f"\nProfile '{profile_id}' already exists. Overwrite? [y/N]: ").strip().lower()
        if answer != "y":
            new_id = input("Enter a different profile ID: ").strip()
            if not new_id:
                print("Aborted.")
                sys.exit(0)
            profile_id = new_id

    # Step 3: Print summary
    domains = extracted.get("domains", {})
    domain_str = ", ".join(f"{k} ({v})" for k, v in sorted(domains.items(), key=lambda x: -x[1]))
    languages = ", ".join(extracted.get("languages", []))
    home_locs = extracted.get("home_locations", [])
    location_str = ", ".join(home_locs[:2]).title() if home_locs else "(unknown)"
    exclude = extracted.get("exclude_companies", [])
    skills = extracted.get("skills", [])
    track = extracted.get("track", "ic")
    current = extracted.get("current_level", "?")
    target = extracted.get("target_level", "?")

    print(f"\n{'─' * 52}")
    print(f"  Name:         {extracted.get('name', '?')}")
    print(f"  Current:      {current.title()} PM ({track.upper()} track)")
    print(f"  Target:       {target.title()} PM")
    print(f"  Domains:      {domain_str}")
    print(f"  Languages:    {languages}")
    print(f"  Location:     {location_str}")
    if exclude:
        shown = ", ".join(exclude[:5])
        suffix = "..." if len(exclude) > 5 else ""
        print(f"  Exclude:      {shown}{suffix}")
    print(f"  Skills:       {len(skills)} extracted")
    print(f"{'─' * 52}\n")

    # Step 4: Interactive questions
    print("Three quick questions:\n")
    salary_min = _ask_salary()
    location_choice, updated_locations = _ask_location(home_locs)
    email = _ask_email(default=extracted.get("email"))

    # Step 5: Generate and save files
    print("\nGenerating profile...")
    profile_yaml = _build_profile_yaml(
        extracted=extracted,
        profile_id=profile_id,
        email=email,
        salary_min=salary_min,
        location_choice=location_choice,
        home_locations=updated_locations,
    )

    out_profile, out_cv, out_seen = _save_files(profile_id, cv_markdown, profile_yaml)

    # Step 6: Print completion
    print(f"\n{'─' * 52}")
    print(f"  Profile saved: {out_profile}")
    print(f"  Knowledge:     {out_cv}")
    print(f"  Seen IDs:      {out_seen}")
    print(f"{'─' * 52}")
    print("\nRun your first scan:")
    print(f"  python main.py --profile {profile_id} --notify\n")


if __name__ == "__main__":
    main()
